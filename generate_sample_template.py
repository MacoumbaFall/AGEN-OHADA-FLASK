from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_template():
    doc = Document()
    
    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Header / Title
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("ÉTUDE DE MAÎTRE {{ current_user.username | upper }}")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph("Notaire à la Résidence de ....................")
    doc.add_paragraph("BP: ........ - Tel: ....................")
    doc.add_paragraph("-" * 20).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Act Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("\nACTE DE TEST - {{ template.type_acte }}")
    run.bold = True
    run.underline = True
    run.font.size = Pt(16)

    # Body
    doc.add_paragraph(f"\nDOSSIER N° : {{{{ dossier.numero_dossier }}}}")
    doc.add_paragraph(f"DATE : {{{{ date }}}}")

    doc.add_paragraph("\nL'an deux mille vingt-cinq,")
    doc.add_paragraph("Le {{ date }},")
    doc.add_paragraph("\nDevant Maître {{ current_user.username }}, Notaire soussigné,")
    
    doc.add_paragraph("\nONT COMPARU :")
    
    # Parties
    p = doc.add_paragraph()
    p.add_run("1 - LE VENDEUR : ").bold = True
    p.add_run("M./Mme {{ vendeur.nom }} {{ vendeur.prenom }}, né(e) le {{ vendeur.date_naissance }}, demeurant à {{ vendeur.adresse }}.")

    p = doc.add_paragraph()
    p.add_run("2 - L'ACHETEUR : ").bold = True
    p.add_run("M./Mme {{ acheteur.nom }} {{ acheteur.prenom }}, né(e) le {{ acheteur.date_naissance }}, demeurant à {{ acheteur.adresse }}.")

    doc.add_paragraph("\nLESQUELS ont requis le Notaire soussigné de dresser l'acte dont la teneur suit :")
    
    # Section
    doc.add_paragraph("\nOBJET DE L'ACTE").bold = True
    doc.add_paragraph("L'objet du présent acte concerne le dossier intitulé : {{ dossier.intitule }}.")
    doc.add_paragraph("Description du dossier : {{ dossier.description }}")

    # Footer/Signatures
    doc.add_paragraph("\n" + "_" * 30)
    doc.add_paragraph("DONT ACTE sur {{ pages | default('...') }} pages.")
    
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3)
    
    cells = table.rows[0].cells
    cells[0].text = "Signature du Vendeur"
    cells[1].text = "Signature de l'Acheteur"

    # Save
    filename = "MODELE_TEST_NOTAIRE.docx"
    doc.save(filename)
    print(f"Template created: {filename}")

if __name__ == "__main__":
    create_test_template()
